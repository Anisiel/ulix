import os
import json
from docx import Document

# =========================================================
# FUNZIONI UTILI
# =========================================================
def format_eur(x):
    return f"{x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

# Bordi tabella
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def aggiungi_bordi(tabella):
    tbl = tabella._element
    tblPr = tbl.xpath("w:tblPr")[0]
    tblBorders = OxmlElement("w:tblBorders")
    for bord in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{bord}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

# =========================================================
# CONFIG
# =========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_JSON = os.path.join(BASE_DIR, "db_unico.json")
INPUT_TXT = os.path.join(BASE_DIR, "lista_report.txt")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "report")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# =========================================================
# CARICA DATABASE
# =========================================================
with open(DB_JSON, "r", encoding="utf-8") as f:
    db = json.load(f)

# =========================================================
# CREA REPORT SINGOLO
# =========================================================
def crea_report(cf, denominazione_txt=None):

    if cf not in db:
        print(f" Codice fiscale {cf} non trovato nel JSON unificato. Saltato.")
        return

    dati_cf = db[cf]

    # Denominazione
    if denominazione_txt and denominazione_txt.strip():
        denominazione = denominazione_txt.strip()
    else:
        # estrai la prima denominazione disponibile
        for cat in dati_cf:
            if len(dati_cf[cat]) > 0:
                denominazione = dati_cf[cat][0]["denominazione"]
                break

    # =========================================================
    # Documento Word
    # =========================================================
    doc = Document()
    doc.add_heading("REPORT SOGGETTO - CREDITI E CONTRIBUTI", level=1)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Denominazione: ").bold = True
    p.add_run(denominazione + "\n")
    p.add_run("Codice Fiscale: ").bold = True
    p.add_run(cf)

    doc.add_paragraph("")

    # =========================================================
    # PER OGNI CATEGORIA
    # =========================================================
    for categoria in sorted(dati_cf.keys()):

        doc.add_heading(f"Categoria: {categoria.upper()}", level=2)

    #    lista = sorted(dati_cf[categoria], key=lambda x: x["anno"])

        lista = sorted(
        [
            r for r in dati_cf[categoria]
            if not (
                r.get("tipo") == "contributi diretti"
                and float(r.get("importo", 0)) <= 0
            )
        ],
        key=lambda x: x["anno"]
    )

        tab = doc.add_table(rows=1, cols=3)
        aggiungi_bordi(tab)
        hdr = tab.rows[0].cells
        hdr[0].text = "Anno"
        hdr[1].text = "Denominazione"
        hdr[2].text = "Importo (€)"

        totale = 0

        for r in lista:
            row = tab.add_row().cells
            row[0].text = str(r["anno"])
            row[1].text = r["denominazione"]

            # Campo dinamico credito / contributo
            valore = r.get("credito", r.get("importo", 0))
            row[2].text = format_eur(valore)

            totale += valore

        doc.add_paragraph("")
        t = doc.add_paragraph()
        t.add_run(f"Totale categoria {categoria}: ").bold = True
        t.add_run(format_eur(totale) + " €")
        doc.add_paragraph("")

    # =========================================================
    # SALVA FILE
    # =========================================================
    nome_file = f"{cf}_{denominazione.replace('/', ' ')}.docx"
    percorso = os.path.join(OUTPUT_FOLDER, nome_file)
    doc.save(percorso)
    print(f" Generato: {nome_file}")


# =========================================================
# CARICA LISTA CF E CREA REPORT
# =========================================================
with open(INPUT_TXT, "r", encoding="utf-8") as f:
    righe = f.readlines()

for riga in righe:
    if ";" not in riga:
        continue

    parti = riga.strip().split(";")
    if len(parti) == 0:
        continue

    cf = parti[0].strip()
    denominazione = parti[1].strip() if len(parti) > 1 else ""

    if cf:
        crea_report(cf, denominazione)