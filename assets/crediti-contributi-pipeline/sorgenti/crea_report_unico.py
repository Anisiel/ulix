import os
import json
from docx import Document

# =========================================================
# UTILS
# =========================================================
def format_eur(x):
    return f"{x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def aggiungi_bordi(tabella):
    tbl = tabella._element
    tblPr = tbl.xpath("w:tblPr")[0]
    tblBorders = OxmlElement("w:tblBorders")
    for bord in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{bord}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

# =========================================================
# CONFIG
# =========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DB_JSON = os.path.join(BASE_DIR, "db_unico.json")
INPUT_TXT = os.path.join(BASE_DIR, "lista_report.txt")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "report")
OUTPUT_FILE = os.path.join(OUTPUT_FOLDER, "Riepilogo_Unico.docx")

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# CARICA DB
with open(DB_JSON, "r", encoding="utf-8") as f:
    db = json.load(f)

# =========================================================
# LEGGI LISTA CF
# =========================================================
lista_cf = []
with open(INPUT_TXT, "r", encoding="utf-8") as f:
    for riga in f.readlines():
        if ";" not in riga:
            continue
        cf = riga.strip().split(";")[0].strip()
        if cf:
            lista_cf.append(cf)

# =========================================================
# CREA REPORT UNICO
# =========================================================
doc = Document()
doc.add_heading("RIEPILOGO GENERALE - CREDITI + CONTRIBUTI", level=1)
doc.add_paragraph("")

for cf in lista_cf:

    if cf not in db:
        p = doc.add_paragraph()
        p.add_run(f" Codice fiscale {cf} non trovato nel JSON.").bold = True
        doc.add_paragraph("")
        continue

    impresa = db[cf]

    # denominazione
    denominazione = None
    for cat in impresa:
        if len(impresa[cat]) > 0:
            denominazione = impresa[cat][0]["denominazione"]
            break

    doc.add_heading(denominazione, level=2)
    doc.add_paragraph(f"Codice Fiscale: {cf}")
    doc.add_paragraph("")

    for categoria in sorted(impresa.keys()):

        doc.add_heading(f"Categoria: {categoria.upper()}", level=3)

#        lista = sorted(impresa[categoria], key=lambda x: x["anno"])
        lista = sorted(
            [
                r for r in impresa[categoria]
                if not (
                    r.get("tipo") == "contributi diretti"
                    and float(r.get("importo", 0)) <= 0
                )
            ],
            key=lambda x: x["anno"]
        )
        totale_cat = sum(r.get("credito", r.get("importo", 0)) for r in lista)

        tab = doc.add_table(rows=1, cols=3)
        aggiungi_bordi(tab)
        hdr = tab.rows[0].cells
        hdr[0].text = "Anno"
        hdr[1].text = "Importo (€)"
        hdr[2].text = "Totale categoria (€)"

        for i, r in enumerate(lista):
            val = r.get("credito", r.get("importo", 0))
            row = tab.add_row().cells
            row[0].text = str(r["anno"])
            row[1].text = format_eur(val)
            row[2].text = format_eur(totale_cat) if i == len(lista) - 1 else ""

        doc.add_paragraph("")
    doc.add_paragraph("\n")

# =========================================================
# SALVA
# =========================================================
doc.save(OUTPUT_FILE)
print("Report unico generato:", OUTPUT_FILE)